from PyQt5.QtGui import QIcon
from PyQt5.QtWidgets import *
from PyQt5 import QtGui
from math import *

import sys
import matplotlib
matplotlib.use('Qt5Agg')
from PyQt5 import QtCore, QtGui, QtWidgets

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_LINE_SPACING

from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT


import pandas as pd
import datetime

import os


class nakliye_hesabi(QWidget):




    def __init__(self):
        super().__init__()
        self.setUI()




    def setUI(self):
        df = pd.read_excel ("pozlar.xlsx")
        # print (df)
        self.k=df.iloc[0]['FİYAT']
        self.poz1=df.iloc[1]['FİYAT']
        self.poz2=df.iloc[2]['FİYAT']
        # print(self.poz1)
        # print(self.poz2)





        self.setWindowTitle("Nakliye Hesabı")
        self.setGeometry(700, 200, 800, 600)
        self.ana=QHBoxLayout()
        self.solsayfa=QHBoxLayout()
        self.sagsayfa=QVBoxLayout()
        self.ic1=QHBoxLayout()

        self.grup1=QGroupBox("Liste")
        self.grup2=QGroupBox("Hesap")

        self.ana.addWidget(self.grup1,40)
        self.ana.addWidget(self.grup2,60)

        self.liste = QListWidget()
        self.liste.setStyleSheet("background : rgb(255,251,186);")
        self.liste.addItem("Çakıl Nakli")
        self.liste.addItem("Kum Nakli")
        self.liste.addItem("Parke,Kaba yonu Taş Nakli")
        self.liste.addItem("Ocak Taşı Nakli")
        self.liste.addItem("Nervürlü Çelik Nakli")
        self.liste.addItem("Çimento Nakli")
        self.liste.addItem("B.A. ve Profil Demiri Nakli")
        self.liste.addItem("Kazı ve Moloz Nakli")

        self.liste.itemClicked.connect(self.on_clicked)

        self.ic1.addWidget(self.liste)
        self.grup1.setLayout(self.ic1)
        #########yükleme-boşaltma###################

        self.pozno=QLabel("Poz :")
        self.fiyat=QLabel("Birim Fiyat :")
        self.miktar=QLabel("Miktar :")

        self.poznotxt=QLineEdit()
        self.poznotxt.setPlaceholderText("10.100.1062")
        self.fiyattxt=QLineEdit()
        self.fiyattxt.setPlaceholderText("Birim Fiyat")
        self.miktartxt=QLineEdit()
        self.miktartxt.setPlaceholderText("Miktar")
        ########################################################


        #########bilgiler###################
        self.Ktxt=QLineEdit()
        self.Ktxt.setText(str(self.k))
        self.Mtxt=QLineEdit()
        self.Mtxt.setPlaceholderText("Mesafe")
        self.Atxt=QLineEdit()
        self.Atxt.setText("1.0")
        self.Gtxt=QLineEdit()
        self.Gtxt.setPlaceholderText("Yoğunluk")
        ########################################################


        #########sonuç###################
        self.Karsiztxt=QLineEdit()
        self.Karsiztxt.setPlaceholderText("Karsız Fiyat)")
        self.kartxt=QLineEdit()
        self.kartxt.setPlaceholderText("Kar")
        self.karlitxt=QLineEdit()
        self.karlitxt.setPlaceholderText("Karlı Fiyat")

        ########################################################
        self.hbox=QHBoxLayout()


        self.hesap=QPushButton("Hesapla")
        self.hesap.clicked.connect(self.Hesapla)
        self.yazdir=QPushButton("Yazdır")
        self.yazdir.clicked.connect(self.Yazdirma)
        self.hbox.addWidget(self.hesap)
        self.hbox.addWidget(self.yazdir)

        self.boş1=QLabel("")
        self.boş2=QLabel("")
        self.boş3=QLabel("")
        self.boş4=QLabel("")
        self.boş5=QLabel("")
        self.boş6=QLabel("")



        form = QFormLayout()

        form.addRow("K :",self.Ktxt)
        form.addRow("M :",self.Mtxt)
        form.addRow("A :",self.Atxt)
        form.addRow("G :",self.Gtxt)
        form.addRow("",self.boş1)


        form.addRow(self.pozno,self.poznotxt)
        form.addRow(self.fiyat,self.fiyattxt)
        form.addRow(self.miktar,self.miktartxt)
        form.addRow("",self.boş2)


        form.addRow("Karsız Fiyat :",self.Karsiztxt)
        form.addRow("Kar :",self.kartxt)
        form.addRow("Karlı Fiyat :",self.karlitxt)
        form.addRow("",self.hbox)

        self.grup2.setLayout(form)



        self.setLayout(self.ana)
        self.show()




    def Hesapla(self):
        try:
            self.K=float(self.Ktxt.text().replace(",", "."))
            self.M=float(self.Mtxt.text().replace(",", "."))+0.0
            self.A=float(self.Atxt.text().replace(",", "."))
            self.G=float(self.Gtxt.text().replace(",", "."))

            #####yükleme-boşaltma########

            # M>10 KM ise F = A x K x (0.0007 x M+0.01) x G
            # M<=10 KM ise F = A x 0.00017 x K x √M x G

            try:
                if self.item_=="Çakıl Nakli" or self.item_=="Kum Nakli" or self.item_=="Parke,Kaba yonu Taş Nakli" or self.item_=="Nervürlü Çelik Nakli" or self.item_=="Çimento Nakli" or self.item_=="B.A. ve Profil Demiri Nakli":
                    self.BF=float(self.fiyattxt.text().replace(",", "."))
                    self.MİKTAR=float(self.miktartxt.text().replace(",", "."))
                    if self.M>10:

                        self.F1=self.A*self.K*(0.0007*self.M+0.01)*self.G
                        self.YB=self.BF*self.MİKTAR
                        self.F=self.F1+self.YB
                        self.kar=self.F*0.25
                        self.ToplamF=self.F+self.kar
                        self.Karsiztxt.setText(str(round(self.F,2)))
                        self.kartxt.setText(str(round(self.kar,2)))
                        self.karlitxt.setText(str(round(self.ToplamF,2)))
                    else:
                        self.F1=self.A*0.00017*self.K*sqrt(self.M*1000)*self.G
                        self.YB=self.BF*self.MİKTAR
                        self.F=self.F1+self.YB
                        self.kar=self.F*0.25
                        self.ToplamF=self.F+self.kar
                        self.Karsiztxt.setText(str(round(self.F,2)))
                        self.kartxt.setText(str(round(self.kar,2)))
                        self.karlitxt.setText(str(round(self.ToplamF,2)))

                elif self.item_=="Ocak Taşı Nakli" or self.item_=="Kazı ve Moloz Nakli":
                    if self.M>10:
                        self.F1=self.A*self.K*(0.0007*self.M+0.01)*self.G
                        self.kar=self.F1*0.25
                        self.ToplamF=self.F1+self.kar
                        self.Karsiztxt.setText(str(round(self.F1,2)))
                        self.kartxt.setText(str(round(self.kar,2)))
                        self.karlitxt.setText(str(round(self.ToplamF,2)))
                    else:
                        self.F1=self.A*0.00017*self.K*sqrt(self.M*1000)*self.G
                        self.kar=self.F1*0.25
                        self.ToplamF=self.F1+self.kar
                        self.Karsiztxt.setText(str(round(self.F1,2)))
                        self.kartxt.setText(str(round(self.kar,2)))
                        self.karlitxt.setText(str(round(self.ToplamF,2)))

                else:
                    pass
            except:
                pass
        except:
            pass

    def on_clicked(self,item):
        self.item_=item.text()

        # self.pozno=QLabel("Poz :")
        # self.fiyat=QLabel("Birim Fiyat :")
        # self.miktar=QLabel("Miktar :")
        #
        # self.poznotxt=QLineEdit()
        # self.poznotxt.setPlaceholderText("10.100.1062")
        # self.fiyattxt=QLineEdit()
        # self.fiyattxt.setPlaceholderText("Birim Fiyat")
        # self.miktartxt=QLineEdit()
        # self.miktartxt.setPlaceholderText("Miktar")


        if self.item_=="Çakıl Nakli":
            self.Gtxt.setText("1.8")
            self.poznotxt.setText("15.100.1062")
            self.fiyattxt.setText(str(self.poz1))
            self.miktartxt.setText("0.75")

            self.poznotxt.setVisible(True)
            self.fiyattxt.setVisible(True)
            self.miktartxt.setVisible(True)

            self.pozno.setVisible(True)
            self.fiyat.setVisible(True)
            self.miktar.setVisible(True)
        elif self.item_=="Kum Nakli":
            self.Gtxt.setText("1.8")
            self.poznotxt.setText("15.100.1062")
            self.fiyattxt.setText(str(self.poz1))
            self.miktartxt.setText("0.75")

            self.poznotxt.setVisible(True)
            self.fiyattxt.setVisible(True)
            self.miktartxt.setVisible(True)

            self.pozno.setVisible(True)
            self.fiyat.setVisible(True)
            self.miktar.setVisible(True)

        elif self.item_=="Parke,Kaba yonu Taş Nakli":
            self.Gtxt.setText("2.0")
            self.poznotxt.setText("15.100.1003")
            self.fiyattxt.setText(str(self.poz2))
            self.miktartxt.setText("1.0")

            self.poznotxt.setVisible(True)
            self.fiyattxt.setVisible(True)
            self.miktartxt.setVisible(True)

            self.pozno.setVisible(True)
            self.fiyat.setVisible(True)
            self.miktar.setVisible(True)
        elif self.item_=="Ocak Taşı Nakli":
            self.Gtxt.setText("1.0")
            self.poznotxt.setText("15.100.1062")
            self.fiyattxt.setText(str(self.poz1))
            self.miktartxt.setText("")
            self.poznotxt.setVisible(False)
            self.fiyattxt.setVisible(False)
            self.miktartxt.setVisible(False)

            self.pozno.setVisible(False)
            self.fiyat.setVisible(False)
            self.miktar.setVisible(False)
        elif self.item_=="Nervürlü Çelik Nakli":
            self.Gtxt.setText("1.0")
            self.poznotxt.setText("15.100.1062")
            self.fiyattxt.setText(str(self.poz1))
            self.miktartxt.setText("1.5")

            self.poznotxt.setVisible(True)
            self.fiyattxt.setVisible(True)
            self.miktartxt.setVisible(True)

            self.pozno.setVisible(True)
            self.fiyat.setVisible(True)
            self.miktar.setVisible(True)
        elif self.item_=="Çimento Nakli":
            self.Gtxt.setText("1.0")
            self.poznotxt.setText("15.100.1003")
            self.fiyattxt.setText(str(self.poz2))
            self.miktartxt.setText("0.5")

            self.poznotxt.setVisible(True)
            self.fiyattxt.setVisible(True)
            self.miktartxt.setVisible(True)

            self.pozno.setVisible(True)
            self.fiyat.setVisible(True)
            self.miktar.setVisible(True)
        elif self.item_=="B.A. ve Profil Demiri Nakli":
            self.Gtxt.setText("1.0")
            self.poznotxt.setText("15.100.1062")
            self.fiyattxt.setText(str(self.poz1))
            self.miktartxt.setText("1.5")

            self.poznotxt.setVisible(True)
            self.fiyattxt.setVisible(True)
            self.miktartxt.setVisible(True)

            self.pozno.setVisible(True)
            self.fiyat.setVisible(True)
            self.miktar.setVisible(True)
        elif self.item_=="Kazı ve Moloz Nakli":
            self.Gtxt.setText("2.0")
            # self.poznotxt.setText("15.100.1062")
            # self.fiyattxt.setText(str(self.poz1))
            # self.miktartxt.setText("")
            self.poznotxt.setVisible(False)
            self.fiyattxt.setVisible(False)
            self.miktartxt.setVisible(False)

            self.pozno.setVisible(False)
            self.fiyat.setVisible(False)
            self.miktar.setVisible(False)
        else:
            pass


    def Yazdirma(self):
        try:
            document = Document()
            logo_path = 'dsi.png'  # Path of the image file
            section = document.sections[0]  # Create a section
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height
            sec_header = section.header  # Create header
            header_tp = sec_header.add_paragraph()  # Add a paragraph in the header, you can add any anything in the paragraph
            header_run = header_tp.add_run()  # Add a run in the paragraph. In the run you can set the values
            header_run.add_picture(logo_path, width=Inches(0.6))  # Add a picture and set width.
            rml_header = "\t\t DSİ 93. ŞUBE MÜDÜRLÜĞÜ - ÖN İNCELEME RAPORU \t"
            header_run.add_text(rml_header)
            header_tp.alignment = 0

            an = datetime.datetime.now()
            yıl = an.year
            idx = self.liste.currentRow()

            document.add_heading('NAKLİYE HESABI', level=2)
            paragraph = document.add_paragraph()
            paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE

            table2 = document.add_table(rows=3, cols=3)
            table2.style = "Table Grid"

            heading_cells = table2.rows[0].cells
            heading_cells[0].text = "Analiz "
            heading_cells[1].text = "NAKLİYE BİRİM FİYAT ANALİZİ"
            heading_cells[2].text = str(yıl)

            heading_cells = table2.rows[1].cells
            heading_cells[0].text = "İş Kalemi/İş Grubu :"
            heading_cells[1].text = "Analiz Adı"
            heading_cells[2].text = "Ölçü Birimi"

            heading_cells = table2.rows[2].cells
            heading_cells[0].text = "NAKLİYE - " + str(idx)
            heading_cells[1].text = str(self.item_)

            if self.item_ == "Çakıl Nakli" or self.item_ == "Kum Nakli" or self.item_ == "Parke,Kaba yonu Taş Nakli" or self.item_ == "Kazı ve Moloz Nakli":
                heading_cells[2].text = "M3"
            elif self.item_ == "Ocak Taşı Nakli" or self.item_ == "Nervürlü Çelik Nakli" or self.item_ == "Çimento Nakli" or self.item_ == "B.A. ve Profil Demiri Nakli":
                heading_cells[2].text = "TON"
            else:
                pass

            table3 = document.add_table(rows=1, cols=1)
            table3.style = "Table Grid"

            heading_cells = table3.rows[0].cells
            heading_cells[0].text = ""

            table1 = document.add_table(rows=7, cols=6)
            table1.style = "Table Grid"

            heading_cells = table1.rows[0].cells
            heading_cells[0].text = "Poz No"
            heading_cells[1].text = "Girdiler"
            heading_cells[2].text = "Birim"
            heading_cells[3].text = "Miktar"
            heading_cells[4].text = "Birim Fiyat"
            heading_cells[5].text = "Tutar"

            heading_cells = table1.rows[1].cells
            heading_cells[0].text = ""
            if self.M > 10:
                heading_cells[1].text = "F = A x K x (0.0007 x M+0.01) x G"
            else:
                heading_cells[1].text = "F = A x 0.00017 x K x √M x G"

            heading_cells[2].text = ""
            heading_cells[3].text = ""
            heading_cells[4].text = str(round(self.F1, 2))
            heading_cells[5].text = str(round(self.F1, 2))

            heading_cells = table1.rows[2].cells
            heading_cells[0].text = "A"
            heading_cells[1].text = "A:Zorluk Katsayısı = " + str(self.A)
            heading_cells[2].text = "AD"
            heading_cells[3].text = ""
            heading_cells[4].text = ""
            heading_cells[5].text = ""

            heading_cells = table1.rows[3].cells
            heading_cells[0].text = "G"
            heading_cells[1].text = "G:Taşınan malzemenin yoğunluğu = " + str(self.G) + " ton/m3"
            heading_cells[2].text = ""
            heading_cells[3].text = ""
            heading_cells[4].text = ""
            heading_cells[5].text = ""

            heading_cells = table1.rows[4].cells
            heading_cells[0].text = "K"
            heading_cells[1].text = "K:Taşıt katsayısı(10.110.1003= " + str(self.Ktxt.text()) + ")"
            heading_cells[2].text = ""
            heading_cells[3].text = ""
            heading_cells[4].text = ""
            heading_cells[5].text = ""

            heading_cells = table1.rows[5].cells
            heading_cells[0].text = "M"
            if self.M > 10:
                heading_cells[1].text = "M:Taşıma mesafesi = " + str(self.M) + " km"
            else:
                heading_cells[1].text = "M:Taşıma mesafesi = " + str(self.M * 1000) + " m"

            heading_cells[2].text = ""
            heading_cells[3].text = ""
            heading_cells[4].text = ""
            heading_cells[5].text = ""

            heading_cells = table1.rows[6].cells
            heading_cells[0].text = self.poznotxt.text()

            if heading_cells[0].text == "15.100.1062":
                heading_cells[1].text = "Düz işçi saatlik ücreti"
            elif heading_cells[0].text == "15.100.1003":
                heading_cells[1].text = "1 m3 her nevi taşın taşıtlara yükleme boşaltma ve figüresi"
            else:
                pass

            heading_cells[2].text = "SA"

            if self.item_ == "Ocak Taşı Nakli" or self.item_ == "Kazı ve Moloz Nakli":
                heading_cells[3].text = ""
                heading_cells[4].text = ""
                heading_cells[5].text = ""

            else:
                heading_cells[3].text = self.miktartxt.text()
                heading_cells[4].text = self.fiyattxt.text()
                heading_cells[5].text = str(round((self.BF * self.MİKTAR), 2))

            table4 = document.add_table(rows=3, cols=2)
            table4.style = "Table Grid"

            if self.item_ == "Ocak Taşı Nakli" or self.item_ == "Kazı ve Moloz Nakli":
                heading_cells = table4.rows[0].cells
                heading_cells[0].text = "Karsız Toplam"
                heading_cells[1].text = str(round(self.F1, 2)) + " TL"
            else:
                heading_cells = table4.rows[0].cells
                heading_cells[0].text = "Karsız Toplam"
                heading_cells[1].text = str(round(self.F, 2)) + " TL"

            heading_cells = table4.rows[1].cells
            heading_cells[0].text = "%25 Kar ve Genel Giderler"
            heading_cells[1].text = str(round(self.kar, 2)) + " TL"

            heading_cells = table4.rows[2].cells
            heading_cells[0].text = "Toplam Tutar"
            heading_cells[1].text = str(round(self.ToplamF, 2)) + " TL"

            col_width_dic = {0: 12, 1: 1}
            for col_num in range(2):
                table4.cell(0, col_num).width = Inches(col_width_dic[col_num])

            # table1.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

            # document.add_page_break()

            document.save('nakliye.docx')

            os.system("start nakliye.docx")
        except:
            pass


if __name__ == "__main__":
    app = QApplication(sys.argv)
    pencere = nakliye_hesabi()
    sys.exit(app.exec())


